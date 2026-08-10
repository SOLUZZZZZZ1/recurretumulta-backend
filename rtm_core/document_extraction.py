"""Extracción documental real para satélites RTM no tráfico.

Este módulo transforma documentos originales en un
``DocumentExtractionPacket``. Su autoridad termina ahí: no valida hechos, no
resuelve familias, no selecciona estrategia, no crea una Previa Jurídica y no
llama a Generate.

El proveedor recibe un esquema cerrado construido desde el catálogo documental.
Los identificadores de expediente y documento, el tipo de fuente y el método de
extracción los inyecta el backend; nunca se aceptan desde la respuesta externa.
"""

from __future__ import annotations

import base64
import io
import json
import mimetypes
import os
import re
from pathlib import PurePosixPath
from typing import Any, Callable, Mapping, Optional, Protocol

import requests
from fastapi import HTTPException
from pydantic import BaseModel, ConfigDict, Field
from pypdf import PdfReader
from docx import Document as DocxDocument

from b2_storage import download_bytes
from rtm_core.document_fact_catalog import (
    DOCUMENT_FACT_CATALOG_VERSION,
    canonical_document_service,
    field_spec,
    minimum_fact_keys,
    registered_fact_keys,
)
from rtm_core.document_normalization import (
    DOCUMENT_EXTRACTION_PACKET_VERSION,
    DocumentExtractionPacket,
    DocumentObservation,
)


SERVICE_DOCUMENT_EXTRACTOR_VERSION = "rtm_service_document_extractor_v1_0"
OPENAI_DOCUMENT_PROVIDER_VERSION = "rtm_openai_responses_document_provider_v1_0"
DETERMINISTIC_DOCUMENT_READER_VERSION = "rtm_deterministic_document_reader_v1_0"

_MAX_DOCUMENTS_DEFAULT = 8
_MAX_DOCUMENT_BYTES_DEFAULT = 8 * 1024 * 1024
_MAX_TOTAL_BYTES_DEFAULT = 20 * 1024 * 1024
_MAX_LOCAL_TEXT_CHARS = 120_000
_MAX_EVIDENCE_CHARS = 500

_TEXT_MIMES = {
    "text/plain",
    "text/csv",
    "text/tab-separated-values",
    "application/json",
    "application/xml",
    "text/xml",
}
_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_PDF_MIME = "application/pdf"
_IMAGE_MIMES = {
    "image/jpeg",
    "image/jpg",
    "image/png",
    "image/webp",
    "image/gif",
    "image/tiff",
    "image/bmp",
}


class _StrictModel(BaseModel):
    model_config = ConfigDict(extra="forbid")


class SourceDocument(_StrictModel):
    id: str = Field(min_length=1)
    case_id: str = Field(min_length=1)
    kind: str = Field(min_length=1)
    mime: str = ""
    b2_bucket: str = Field(min_length=1)
    b2_key: str = Field(min_length=1)
    size_bytes: int = Field(default=0, ge=0)
    sha256: Optional[str] = None

    @property
    def filename(self) -> str:
        value = PurePosixPath(self.b2_key).name
        return value or f"{self.id}.bin"


class ProviderObservation(_StrictModel):
    field: str = Field(min_length=1)
    value: Any = None
    page_index: Optional[int] = Field(default=None, ge=0)
    evidence: Optional[str] = None
    confidence: float = Field(ge=0.0, le=1.0)
    notes: list[str] = Field(default_factory=list)


class ProviderDocumentResult(_StrictModel):
    observations: list[ProviderObservation] = Field(default_factory=list)
    unresolved_fields: list[str] = Field(default_factory=list)
    quality_flags: list[str] = Field(default_factory=list)
    document_notes: list[str] = Field(default_factory=list)


class DocumentExtractionDiagnostic(_StrictModel):
    document_id: str
    filename: str
    mime: str
    input_mode: str
    provider_version: str
    model: str
    observation_count: int = 0
    unresolved_count: int = 0
    quality_flags: list[str] = Field(default_factory=list)
    notes: list[str] = Field(default_factory=list)
    error: Optional[str] = None


class ServiceDocumentExtractionResult(_StrictModel):
    authority: str = "rtm_service_document_extractor"
    version: str = SERVICE_DOCUMENT_EXTRACTOR_VERSION
    packet_version: str = DOCUMENT_EXTRACTION_PACKET_VERSION
    catalog_version: str = DOCUMENT_FACT_CATALOG_VERSION
    case_id: str
    service: str
    provider_version: str
    model: str
    packet: DocumentExtractionPacket
    diagnostics: list[DocumentExtractionDiagnostic] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list)


class DocumentProvider(Protocol):
    version: str
    model: str

    def extract_document(
        self,
        *,
        service: str,
        document: SourceDocument,
        content: bytes,
    ) -> tuple[ProviderDocumentResult, str, list[str]]:
        """Devuelve resultado, modo de entrada y avisos locales."""


def _int_env(name: str, default: int, *, minimum: int = 1) -> int:
    raw = (os.getenv(name) or "").strip()
    if not raw:
        return default
    try:
        return max(minimum, int(raw))
    except Exception:
        return default


def extraction_limits() -> dict[str, int]:
    return {
        "max_documents": _int_env(
            "RTM_DOCUMENT_MAX_FILES",
            _MAX_DOCUMENTS_DEFAULT,
        ),
        "max_document_bytes": _int_env(
            "RTM_DOCUMENT_MAX_BYTES",
            _MAX_DOCUMENT_BYTES_DEFAULT,
        ),
        "max_total_bytes": _int_env(
            "RTM_DOCUMENT_MAX_TOTAL_BYTES",
            _MAX_TOTAL_BYTES_DEFAULT,
        ),
    }


def _clean_text(value: Any, *, limit: int) -> str:
    text_value = re.sub(
        r"\s+",
        " ",
        str(value or "").replace("\x00", " ").replace("\u00a0", " "),
    ).strip()
    if len(text_value) > limit:
        return text_value[: limit - 3].rstrip() + "..."
    return text_value


def _normalise_mime(document: SourceDocument) -> str:
    raw = str(document.mime or "").split(";", 1)[0].strip().lower()
    if raw:
        if raw == "image/jpg":
            return "image/jpeg"
        return raw

    guessed, _ = mimetypes.guess_type(document.filename)
    if guessed:
        return guessed.lower()

    suffix = PurePosixPath(document.filename).suffix.lower()
    by_suffix = {
        ".pdf": _PDF_MIME,
        ".docx": _DOCX_MIME,
        ".txt": "text/plain",
        ".csv": "text/csv",
        ".json": "application/json",
        ".xml": "application/xml",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".webp": "image/webp",
        ".gif": "image/gif",
        ".tif": "image/tiff",
        ".tiff": "image/tiff",
        ".bmp": "image/bmp",
    }
    return by_suffix.get(suffix, "application/octet-stream")


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _docx_text(content: bytes) -> str:
    try:
        document = DocxDocument(io.BytesIO(content))
    except Exception as exc:
        raise HTTPException(
            status_code=422,
            detail=f"No puede leerse el DOCX: {type(exc).__name__}",
        ) from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            parts.append(value)

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"[[TABLA {table_index}]]")
        for row in table.rows:
            cells = [
                re.sub(r"\s+", " ", cell.text).strip()
                for cell in row.cells
            ]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def _pdf_page_count(content: bytes) -> Optional[int]:
    try:
        return len(PdfReader(io.BytesIO(content)).pages)
    except Exception:
        return None


def _data_url(mime: str, content: bytes) -> str:
    encoded = base64.b64encode(content).decode("ascii")
    return f"data:{mime};base64,{encoded}"


def _convert_tiff_to_png(content: bytes) -> tuple[bytes, str, list[str]]:
    try:
        from PIL import Image
    except Exception as exc:
        raise HTTPException(
            status_code=422,
            detail=(
                "La imagen TIFF requiere Pillow para convertirse antes de "
                "la extracción."
            ),
        ) from exc

    try:
        image = Image.open(io.BytesIO(content))
        image.seek(0)
        image = image.convert("RGB")
        output = io.BytesIO()
        image.save(output, format="PNG", optimize=True)
        return (
            output.getvalue(),
            "image/png",
            ["tiff_first_frame_converted_to_png"],
        )
    except Exception as exc:
        raise HTTPException(
            status_code=422,
            detail=f"No puede convertirse la imagen TIFF: {type(exc).__name__}",
        ) from exc


def document_response_schema(service: str) -> dict[str, Any]:
    canonical = canonical_document_service(service)
    keys = list(registered_fact_keys(canonical))
    field_schema: dict[str, Any] = {
        "type": "string",
        "enum": keys,
    }
    nullable_value: dict[str, Any] = {
        "anyOf": [
            {"type": "string"},
            {"type": "number"},
            {"type": "integer"},
            {"type": "boolean"},
            {"type": "null"},
        ]
    }
    return {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "observations": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "field": field_schema,
                        "value": nullable_value,
                        "page_index": {
                            "anyOf": [
                                {"type": "integer", "minimum": 0},
                                {"type": "null"},
                            ]
                        },
                        "evidence": {
                            "anyOf": [
                                {"type": "string"},
                                {"type": "null"},
                            ]
                        },
                        "confidence": {
                            "type": "number",
                            "minimum": 0,
                            "maximum": 1,
                        },
                        "notes": {
                            "type": "array",
                            "items": {"type": "string"},
                        },
                    },
                    "required": [
                        "field",
                        "value",
                        "page_index",
                        "evidence",
                        "confidence",
                        "notes",
                    ],
                },
            },
            "unresolved_fields": {
                "type": "array",
                "items": field_schema,
            },
            "quality_flags": {
                "type": "array",
                "items": {"type": "string"},
            },
            "document_notes": {
                "type": "array",
                "items": {"type": "string"},
            },
        },
        "required": [
            "observations",
            "unresolved_fields",
            "quality_flags",
            "document_notes",
        ],
    }


def _field_lines(service: str) -> str:
    values: list[str] = []
    for key in registered_fact_keys(service):
        spec = field_spec(service, key)
        if spec is None:
            continue
        values.append(
            f"- {spec.key}: {spec.label}; tipo={spec.value_type}; "
            f"repetición={spec.merge_mode}"
        )
    return "\n".join(values)


def _provider_instructions(service: str) -> str:
    minimum = ", ".join(minimum_fact_keys(service)) or "ninguno"
    return (
        "Actúas exclusivamente como lector documental de RTM. "
        "Extrae solo datos que se vean literalmente en el documento. "
        "No clasifiques el asunto, no elijas familia, no propongas estrategia, "
        "no redactes argumentos, no calcules plazos jurídicos y no determines "
        "qué reclamación procede.\n\n"
        "Reglas obligatorias:\n"
        "1. Cada observación debe incluir un fragmento breve de evidencia del "
        "propio documento y confianza de 0 a 1.\n"
        "2. page_index es cero-based cuando la página pueda determinarse; "
        "si no puede, usa null.\n"
        "3. No copies el documento completo. La evidencia debe ser el mínimo "
        "fragmento que sostiene el dato.\n"
        "4. Una duda, dato manuscrito inseguro, lectura incompleta o valor "
        "contradictorio se declara en unresolved_fields o quality_flags; "
        "nunca se completa por probabilidad.\n"
        "5. No extraigas DNI, NIE, domicilio, firma, datos de pago o "
        "autorización, salvo que el catálogo cerrado los admita expresamente.\n"
        "6. No devuelvas OCR crudo, resúmenes jurídicos, familia, scoring, "
        "estrategia, borrador ni recomendaciones.\n"
        f"7. Servicio del expediente: {service}.\n"
        f"8. Campos mínimos de orientación: {minimum}.\n\n"
        "Campos permitidos:\n"
        f"{_field_lines(service)}"
    )


def _text_input(
    *,
    document: SourceDocument,
    content: bytes,
    mime: str,
) -> tuple[list[dict[str, Any]], str, list[str]]:
    warnings: list[str] = []
    if mime == _DOCX_MIME:
        text_value = _docx_text(content)
    else:
        text_value = _decode_text(content)

    if not text_value.strip():
        raise HTTPException(
            status_code=422,
            detail="El documento de texto no contiene contenido legible.",
        )
    if len(text_value) > _MAX_LOCAL_TEXT_CHARS:
        text_value = text_value[:_MAX_LOCAL_TEXT_CHARS]
        warnings.append("local_text_truncated")

    return (
        [
            {
                "type": "input_text",
                "text": (
                    f"Documento: {document.filename}\n"
                    "Contenido documental:\n"
                    f"{text_value}"
                ),
            }
        ],
        "document_text",
        warnings,
    )


def _provider_content(
    *,
    document: SourceDocument,
    content: bytes,
) -> tuple[list[dict[str, Any]], str, str, list[str]]:
    mime = _normalise_mime(document)
    warnings: list[str] = []

    if mime in _TEXT_MIMES or mime == _DOCX_MIME:
        items, mode, warnings = _text_input(
            document=document,
            content=content,
            mime=mime,
        )
        return items, mode, mime, warnings

    if mime == _PDF_MIME:
        page_count = _pdf_page_count(content)
        if page_count is not None:
            warnings.append(f"pdf_pages:{page_count}")
        return (
            [
                {
                    "type": "input_text",
                    "text": (
                        "Lee el PDF adjunto conforme al esquema cerrado. "
                        "No devuelvas una transcripción completa."
                    ),
                },
                {
                    "type": "input_file",
                    "filename": document.filename,
                    "file_data": _data_url(mime, content),
                },
            ],
            "document_vision",
            mime,
            warnings,
        )

    if mime in _IMAGE_MIMES:
        image_content = content
        image_mime = mime
        if mime == "image/tiff":
            image_content, image_mime, converted = _convert_tiff_to_png(content)
            warnings.extend(converted)
        return (
            [
                {
                    "type": "input_text",
                    "text": (
                        "Lee la imagen adjunta conforme al esquema cerrado. "
                        "Si algo no es legible, no lo completes."
                    ),
                },
                {
                    "type": "input_image",
                    "image_url": _data_url(image_mime, image_content),
                    "detail": "high",
                },
            ],
            "document_vision",
            image_mime,
            warnings,
        )

    raise HTTPException(
        status_code=415,
        detail={
            "message": "Formato documental no soportado por el extractor.",
            "mime": mime,
            "filename": document.filename,
        },
    )


def build_responses_payload(
    *,
    service: str,
    document: SourceDocument,
    content: bytes,
    model: str,
) -> tuple[dict[str, Any], str, str, list[str]]:
    canonical = canonical_document_service(service)
    items, mode, effective_mime, warnings = _provider_content(
        document=document,
        content=content,
    )
    payload = {
        "model": model,
        "store": False,
        "input": [
            {
                "role": "developer",
                "content": [
                    {
                        "type": "input_text",
                        "text": _provider_instructions(canonical),
                    }
                ],
            },
            {
                "role": "user",
                "content": items,
            },
        ],
        "text": {
            "format": {
                "type": "json_schema",
                "name": "rtm_document_extraction",
                "strict": True,
                "schema": document_response_schema(canonical),
            }
        },
    }
    return payload, mode, effective_mime, warnings


def _response_output_text(data: Mapping[str, Any]) -> str:
    direct = data.get("output_text")
    if isinstance(direct, str) and direct.strip():
        return direct.strip()

    parts: list[str] = []
    output = data.get("output")
    if isinstance(output, list):
        for item in output:
            if not isinstance(item, Mapping):
                continue
            content = item.get("content")
            if not isinstance(content, list):
                continue
            for entry in content:
                if not isinstance(entry, Mapping):
                    continue
                if entry.get("type") in {"output_text", "text"}:
                    value = entry.get("text")
                    if isinstance(value, str):
                        parts.append(value)
    return "".join(parts).strip()


def parse_provider_response(
    data: Mapping[str, Any],
    *,
    service: str,
) -> ProviderDocumentResult:
    canonical = canonical_document_service(service)
    raw = _response_output_text(data)
    if not raw:
        raise HTTPException(
            status_code=502,
            detail="El proveedor documental no devolvió contenido estructurado.",
        )
    try:
        payload = json.loads(raw)
    except Exception as exc:
        raise HTTPException(
            status_code=502,
            detail="El proveedor documental devolvió JSON inválido.",
        ) from exc
    if not isinstance(payload, Mapping):
        raise HTTPException(
            status_code=502,
            detail="La respuesta documental no es un objeto JSON.",
        )

    observations: list[ProviderObservation] = []
    raw_observations = payload.get("observations")
    if isinstance(raw_observations, list):
        for item in raw_observations:
            if not isinstance(item, Mapping):
                continue
            field = str(item.get("field") or "").strip()
            spec = field_spec(canonical, field)
            if spec is None:
                continue
            evidence = _clean_text(
                item.get("evidence"),
                limit=_MAX_EVIDENCE_CHARS,
            )
            notes = item.get("notes")
            if not isinstance(notes, list):
                notes = []
            try:
                observations.append(
                    ProviderObservation(
                        field=spec.key,
                        value=item.get("value"),
                        page_index=item.get("page_index"),
                        evidence=evidence or None,
                        confidence=float(item.get("confidence") or 0.0),
                        notes=[
                            _clean_text(note, limit=300)
                            for note in notes
                            if _clean_text(note, limit=300)
                        ][:8],
                    )
                )
            except Exception:
                continue

    unresolved: list[str] = []
    raw_unresolved = payload.get("unresolved_fields")
    if isinstance(raw_unresolved, list):
        for item in raw_unresolved:
            spec = field_spec(canonical, str(item or ""))
            if spec is not None and spec.key not in unresolved:
                unresolved.append(spec.key)

    quality = payload.get("quality_flags")
    quality_flags = (
        [
            _clean_text(item, limit=120)
            for item in quality
            if _clean_text(item, limit=120)
        ][:20]
        if isinstance(quality, list)
        else []
    )
    notes = payload.get("document_notes")
    document_notes = (
        [
            _clean_text(item, limit=500)
            for item in notes
            if _clean_text(item, limit=500)
        ][:20]
        if isinstance(notes, list)
        else []
    )
    return ProviderDocumentResult(
        observations=observations,
        unresolved_fields=unresolved,
        quality_flags=quality_flags,
        document_notes=document_notes,
    )


class OpenAIResponsesDocumentProvider:
    version = OPENAI_DOCUMENT_PROVIDER_VERSION

    def __init__(
        self,
        *,
        api_key: Optional[str] = None,
        model: Optional[str] = None,
        timeout_seconds: Optional[int] = None,
    ) -> None:
        self._api_key = api_key
        self.model = (
            str(model or "").strip()
            or (os.getenv("OPENAI_DOCUMENT_MODEL") or "").strip()
            or (os.getenv("OPENAI_MODEL") or "").strip()
            or "gpt-4o"
        )
        self.timeout_seconds = (
            int(timeout_seconds)
            if timeout_seconds is not None
            else _int_env("OPENAI_DOCUMENT_TIMEOUT_SECONDS", 120)
        )

    @property
    def api_key(self) -> str:
        value = (
            str(self._api_key or "").strip()
            or (os.getenv("OPENAI_API_KEY") or "").strip()
        )
        if not value:
            raise HTTPException(
                status_code=500,
                detail="OPENAI_API_KEY no configurado para extracción documental.",
            )
        return value

    def extract_document(
        self,
        *,
        service: str,
        document: SourceDocument,
        content: bytes,
    ) -> tuple[ProviderDocumentResult, str, list[str]]:
        payload, mode, _effective_mime, warnings = build_responses_payload(
            service=service,
            document=document,
            content=content,
            model=self.model,
        )
        try:
            response = requests.post(
                "https://api.openai.com/v1/responses",
                headers={
                    "Authorization": f"Bearer {self.api_key}",
                    "Content-Type": "application/json",
                },
                json=payload,
                timeout=self.timeout_seconds,
            )
        except requests.RequestException as exc:
            raise HTTPException(
                status_code=502,
                detail=(
                    "No ha podido contactarse con el proveedor de extracción "
                    f"documental: {type(exc).__name__}"
                ),
            ) from exc

        if not response.ok:
            detail = _clean_text(response.text, limit=500)
            raise HTTPException(
                status_code=502,
                detail={
                    "message": "El proveedor de extracción documental devolvió error.",
                    "status_code": response.status_code,
                    "provider_detail": detail,
                },
            )
        try:
            data = response.json()
        except Exception as exc:
            raise HTTPException(
                status_code=502,
                detail="Respuesta no JSON del proveedor documental.",
            ) from exc

        return (
            parse_provider_response(data, service=service),
            mode,
            warnings,
        )


def get_document_provider() -> DocumentProvider:
    return OpenAIResponsesDocumentProvider()


def _diagnostic_error(
    document: SourceDocument,
    *,
    provider: DocumentProvider,
    mime: str,
    exc: Exception,
) -> DocumentExtractionDiagnostic:
    if isinstance(exc, HTTPException):
        detail = exc.detail
    else:
        detail = f"{type(exc).__name__}: {exc}"
    if not isinstance(detail, str):
        detail = json.dumps(detail, ensure_ascii=False, default=str)
    return DocumentExtractionDiagnostic(
        document_id=document.id,
        filename=document.filename,
        mime=mime,
        input_mode="failed",
        provider_version=provider.version,
        model=provider.model,
        error=_clean_text(detail, limit=700),
    )


def extract_service_documents(
    *,
    case_id: str,
    service: str,
    documents: list[SourceDocument],
    provider: Optional[DocumentProvider] = None,
    byte_loader: Optional[Callable[[str, str], bytes]] = None,
) -> ServiceDocumentExtractionResult:
    canonical = canonical_document_service(service)
    if not documents:
        raise HTTPException(
            status_code=409,
            detail="No hay documentos seleccionados para la extracción.",
        )
    limits = extraction_limits()
    if len(documents) > limits["max_documents"]:
        raise HTTPException(
            status_code=413,
            detail={
                "message": "Se supera el número máximo de documentos por ejecución.",
                "maximum": limits["max_documents"],
                "received": len(documents),
            },
        )
    if any(document.case_id != case_id for document in documents):
        raise HTTPException(
            status_code=409,
            detail="La selección contiene documentos de otro expediente.",
        )

    declared_total = sum(max(0, int(document.size_bytes or 0)) for document in documents)
    if declared_total > limits["max_total_bytes"]:
        raise HTTPException(
            status_code=413,
            detail={
                "message": "El conjunto documental supera el límite de tamaño.",
                "maximum_bytes": limits["max_total_bytes"],
                "received_bytes": declared_total,
            },
        )

    selected_provider = provider or get_document_provider()
    loader = byte_loader or download_bytes
    observations: list[DocumentObservation] = []
    unresolved: set[str] = set()
    quality_flags: set[str] = set()
    diagnostics: list[DocumentExtractionDiagnostic] = []
    warnings: list[str] = []
    successful_documents: list[str] = []
    actual_total = 0

    for document in documents:
        mime = _normalise_mime(document)
        if (
            document.size_bytes
            and document.size_bytes > limits["max_document_bytes"]
        ):
            raise HTTPException(
                status_code=413,
                detail={
                    "message": "Un documento supera el límite individual.",
                    "document_id": document.id,
                    "maximum_bytes": limits["max_document_bytes"],
                    "received_bytes": document.size_bytes,
                },
            )
        try:
            content = loader(document.b2_bucket, document.b2_key)
            if not content:
                raise HTTPException(
                    status_code=422,
                    detail="El objeto almacenado está vacío.",
                )
            actual_total += len(content)
            if len(content) > limits["max_document_bytes"]:
                raise HTTPException(
                    status_code=413,
                    detail={
                        "message": "Un documento supera el límite individual.",
                        "document_id": document.id,
                        "maximum_bytes": limits["max_document_bytes"],
                        "received_bytes": len(content),
                    },
                )
            if actual_total > limits["max_total_bytes"]:
                raise HTTPException(
                    status_code=413,
                    detail={
                        "message": "El conjunto descargado supera el límite total.",
                        "maximum_bytes": limits["max_total_bytes"],
                        "received_bytes": actual_total,
                    },
                )

            provider_result, input_mode, local_warnings = (
                selected_provider.extract_document(
                    service=canonical,
                    document=document,
                    content=content,
                )
            )
            source_type = (
                "document_text"
                if input_mode == "document_text"
                else "document_vision"
            )
            extraction_method = (
                f"{SERVICE_DOCUMENT_EXTRACTOR_VERSION}+"
                f"{selected_provider.version}+{selected_provider.model}"
            )
            for item in provider_result.observations:
                spec = field_spec(canonical, item.field)
                if spec is None:
                    continue
                observations.append(
                    DocumentObservation(
                        field=spec.key,
                        value=item.value,
                        document_id=document.id,
                        page_index=item.page_index,
                        evidence=item.evidence,
                        confidence=item.confidence,
                        extraction_method=extraction_method,
                        source_type=source_type,
                        notes=list(item.notes),
                    )
                )
            unresolved.update(provider_result.unresolved_fields)
            quality_flags.update(provider_result.quality_flags)
            quality_flags.update(local_warnings)
            successful_documents.append(document.id)
            diagnostics.append(
                DocumentExtractionDiagnostic(
                    document_id=document.id,
                    filename=document.filename,
                    mime=mime,
                    input_mode=input_mode,
                    provider_version=selected_provider.version,
                    model=selected_provider.model,
                    observation_count=len(provider_result.observations),
                    unresolved_count=len(provider_result.unresolved_fields),
                    quality_flags=sorted(
                        set(provider_result.quality_flags + local_warnings)
                    ),
                    notes=provider_result.document_notes,
                )
            )
        except HTTPException as exc:
            if exc.status_code == 413:
                raise
            diagnostics.append(
                _diagnostic_error(
                    document,
                    provider=selected_provider,
                    mime=mime,
                    exc=exc,
                )
            )
            quality_flags.add("partial_extraction_failure")
        except Exception as exc:
            diagnostics.append(
                _diagnostic_error(
                    document,
                    provider=selected_provider,
                    mime=mime,
                    exc=exc,
                )
            )
            quality_flags.add("partial_extraction_failure")

    if not successful_documents:
        raise HTTPException(
            status_code=502,
            detail={
                "message": "No ha podido extraerse ningún documento.",
                "diagnostics": [
                    diagnostic.model_dump(mode="json")
                    for diagnostic in diagnostics
                ],
            },
        )

    observed_keys = {item.field for item in observations}
    for key in minimum_fact_keys(canonical):
        if key not in observed_keys:
            unresolved.add(key)

    if len(successful_documents) < len(documents):
        warnings.append(
            "La ejecución es parcial: uno o más documentos requieren revisión."
        )
    if not observations:
        warnings.append(
            "El proveedor no produjo observaciones consolidables; OPS debe revisar."
        )

    packet = DocumentExtractionPacket(
        case_id=case_id,
        service=canonical,
        extractor_version=(
            f"{SERVICE_DOCUMENT_EXTRACTOR_VERSION}+"
            f"{selected_provider.version}+{selected_provider.model}"
        ),
        source_document_ids=[document.id for document in documents],
        observations=observations,
        declared_unresolved=sorted(unresolved),
        quality_flags=sorted(quality_flags),
    )
    return ServiceDocumentExtractionResult(
        case_id=case_id,
        service=canonical,
        provider_version=selected_provider.version,
        model=selected_provider.model,
        packet=packet,
        diagnostics=diagnostics,
        warnings=warnings,
    )
