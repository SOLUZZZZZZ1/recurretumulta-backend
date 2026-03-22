import io
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SECTION_TITLES = {
    "ANTECEDENTES",
    "ALEGACIONES",
    "FUNDAMENTOS DE DERECHO",
    "SUPLICA",
    "S U P L I C A",
    "OTROSÍ DIGO",
    "OTROSI DIGO",
}

def _iter_bold_segments(text: str):
    parts = re.split(r'(\*\*.*?\*\*)', text or "")
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            yield part[2:-2], True
        else:
            yield part, False

def _write_markdown_bold(paragraph, text: str):
    for chunk, is_bold in _iter_bold_segments(text):
        run = paragraph.add_run(chunk)
        run.bold = is_bold

def build_docx(title: str, body: str) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for raw_line in (body or "").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

        line = raw_line.rstrip()
        txt = line.strip().upper()

        if not line.strip():
            continue

        if "ESCRITO DE ALEGACIONES" in txt:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            continue

        if txt.startswith("A LA "):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _write_markdown_bold(p, line)
            continue

        if txt in SECTION_TITLES or txt.startswith("I. ") or txt.startswith("II. ") or txt.startswith("III. ") or txt.startswith("ALEGACIÓN "):
            run = p.add_run(line.replace("**", ""))
            run.bold = True
            p.paragraph_format.space_before = Pt(8)
            continue

        _write_markdown_bold(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
